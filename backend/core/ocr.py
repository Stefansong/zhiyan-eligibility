"""本地文档识别：照片 / PDF / Excel → 逐页文字。

图片不发送到任何外部服务器，OCR 全部在本机完成（RapidOCR）。
Excel 用 openpyxl 直接结构化读取，不走 OCR。
"""
import io
from pathlib import Path

import fitz  # PyMuPDF
import numpy as np
from PIL import Image, ImageOps

_engine = None

# 文本型 PDF 判定阈值：一页抽出的字符数超过该值则认为有文本层
_TEXT_LAYER_MIN_CHARS = 50

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".heic", ".bmp", ".webp", ".tif", ".tiff"}
EXCEL_EXTS = {".xlsx", ".xlsm", ".xls"}


def get_engine():
    global _engine
    if _engine is None:
        from rapidocr import RapidOCR
        _engine = RapidOCR()
    return _engine


def _ocr_ndarray(img: np.ndarray) -> str:
    result = get_engine()(img)
    if result is None or result.txts is None:
        return ""
    return "\n".join(result.txts)


def image_bytes_to_text(data: bytes) -> str:
    """单张照片 → 文字。自动按 EXIF 方向纠正。"""
    img = Image.open(io.BytesIO(data))
    img = ImageOps.exif_transpose(img).convert("RGB")
    # 过大的照片缩到长边 2500px，兼顾速度与识别率
    if max(img.size) > 2500:
        img.thumbnail((2500, 2500))
    return _ocr_ndarray(np.array(img))


def pdf_bytes_to_pages(data: bytes) -> list[dict]:
    """PDF → [{page, text, source}]。有文本层直接抽取，扫描页渲染后 OCR。"""
    pages = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for i, page in enumerate(doc, start=1):
            text = page.get_text().strip()
            if len(text) >= _TEXT_LAYER_MIN_CHARS:
                pages.append({"page": i, "text": text, "source": "文本层"})
            else:
                pix = page.get_pixmap(dpi=300)
                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
                if pix.n == 4:
                    img = img[:, :, :3]
                pages.append({"page": i, "text": _ocr_ndarray(img), "source": "OCR"})
    return pages


def docx_bytes_to_text(data: bytes) -> str:
    """Word 文档 → 文字（python-docx）。"""
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def xlsx_bytes_to_grid(data: bytes, max_rows: int = 200, max_cols: int = 20) -> dict:
    """Excel → {cols, rows, text}。直接结构化读取（首个有数据的工作表）。"""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(max_row=max_rows, max_col=max_cols, values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                rows.append(cells)
        if rows:
            width = max(len(r) for r in rows)
            rows = [r + [""] * (width - len(r)) for r in rows]
            cols, body = rows[0], rows[1:]
            text = "\n".join("\t".join(r) for r in rows)
            return {"sheet": ws.title, "cols": cols, "rows": body, "text": text}
    return {"sheet": "", "cols": [], "rows": [], "text": ""}


def file_to_pages(filename: str, data: bytes) -> tuple[list[dict], dict | None]:
    """任意上传文件 → (逐页文字列表, Excel 网格或 None)。"""
    suffix = Path(filename).suffix.lower()
    grid = None
    if suffix == ".pdf":
        pages = pdf_bytes_to_pages(data)
    elif suffix in EXCEL_EXTS:
        grid = xlsx_bytes_to_grid(data)
        pages = [{"page": 1, "text": grid["text"], "source": "数据表导入"}]
    elif suffix == ".docx":
        pages = [{"page": 1, "text": docx_bytes_to_text(data), "source": "文本层"}]
    else:
        pages = [{"page": 1, "text": image_bytes_to_text(data), "source": "OCR"}]
    for p in pages:
        p["file"] = filename
    return pages, grid


def pages_to_text(pages: list[dict]) -> str:
    """逐页文字合并为带来源标记的全文。"""
    parts = []
    for p in pages:
        parts.append(f"【文件：{p['file']} 第{p['page']}页】\n{p['text']}")
    return "\n\n".join(parts)
